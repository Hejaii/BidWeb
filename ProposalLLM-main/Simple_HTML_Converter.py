import os
import re
import requests
from docx import Document
from docx.shared import Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
import tempfile
from PIL import Image, ImageDraw, ImageFont
import io
import html2image

class SimpleHTMLConverter:
    def __init__(self):
        # 通义千问 API Key
        self.api_key = "sk-fe0485c281964259b404907d483d3777"
        self.api_url = "https://dashscope.aliyuncs.com/api/v1/services/aigc/text-generation/generation"
        
    def call_qwen_api(self, prompt):
        """调用通义千问API"""
        headers = {
            "Authorization": f"Bearer {self.api_key}",
            "Content-Type": "application/json"
        }
        payload = {
            "model": "qwen-turbo",
            "input": {
                "messages": [
                    {
                        "role": "user",
                        "content": prompt
                    }
                ]
            }
        }
        response = requests.post(self.api_url, headers=headers, json=payload)
        return response.json()
    
    def html_to_image(self, html_content):
        """将HTML内容转换为图片"""
        try:
            # 使用html2image将HTML转换为图片
            hti = html2image.Html2Image()
            
            # 创建完整的HTML文档
            full_html = f"""
            <!DOCTYPE html>
            <html>
            <head>
                <meta charset="UTF-8">
                <style>
                    body {{ 
                        font-family: 'Microsoft YaHei', Arial, sans-serif; 
                        margin: 0; 
                        padding: 20px; 
                        background: white; 
                        color: #333;
                    }}
                    .container {{ 
                        max-width: 600px; 
                        margin: 0 auto; 
                        background: white;
                        border: 1px solid #ddd;
                        border-radius: 8px;
                        padding: 20px;
                        box-shadow: 0 2px 4px rgba(0,0,0,0.1);
                    }}
                    .tech-box {{
                        background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
                        color: white;
                        padding: 15px;
                        border-radius: 8px;
                        margin: 10px 0;
                        text-align: center;
                    }}
                    .architecture {{
                        display: flex;
                        justify-content: space-between;
                        margin: 20px 0;
                    }}
                    .layer {{
                        background: rgba(255,255,255,0.2);
                        padding: 10px;
                        border-radius: 5px;
                        text-align: center;
                        flex: 1;
                        margin: 0 5px;
                    }}
                </style>
            </head>
            <body>
                <div class="container">
                    {html_content}
                </div>
            </body>
            </html>
            """
            
            # 生成图片 - 使用正确的API
            output_path = hti.screenshot(html_str=full_html, save_as='temp_image.png')[0]
            
            # 读取生成的图片
            with open(output_path, 'rb') as f:
                image_data = f.read()
            
            # 清理临时文件
            os.remove(output_path)
            
            return image_data
            
        except Exception as e:
            print(f"HTML转图片时出错: {e}")
            return self.create_default_image()
    
    def create_default_image(self):
        """创建默认图片"""
        try:
            img = Image.new('RGB', (600, 400), color='#ecf0f1')
            draw = ImageDraw.Draw(img)
            
            try:
                font = ImageFont.truetype("arial.ttf", 20)
            except:
                font = ImageFont.load_default()
            
            draw.text((200, 190), "技术架构图", fill='#7f8c8d', font=font)
            
            img_buffer = io.BytesIO()
            img.save(img_buffer, format='PNG')
            img_buffer.seek(0)
            
            return img_buffer.getvalue()
        except Exception as e:
            print(f"创建默认图片时出错: {e}")
            return b""
    
    def add_image_to_docx(self, doc, image_data, width_cm=14.0):
        """将图片添加到docx文档"""
        try:
            # 将图片数据写入临时文件
            with tempfile.NamedTemporaryFile(delete=False, suffix='.png') as temp_file:
                temp_file.write(image_data)
                temp_file_path = temp_file.name
            
            # 添加图片到文档
            paragraph = doc.add_paragraph()
            run = paragraph.add_run()
            run.add_picture(temp_file_path, width=Cm(width_cm))
            
            # 设置图片居中
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # 清理临时文件
            os.unlink(temp_file_path)
            
        except Exception as e:
            print(f"添加图片到文档时出错: {e}")
    
    def clean_text_content(self, text):
        """清理文本内容，移除HTML标签"""
        # 移除HTML标签
        clean_text = re.sub(r'<[^>]+>', '', text)
        # 清理多余的空行
        clean_text = re.sub(r'\n\s*\n', '\n\n', clean_text)
        return clean_text.strip()
    
    def convert_txt_to_docx(self, txt_file_path, output_docx_path):
        """将带HTML的txt文件转换为docx文件"""
        print(f"开始转换文件: {txt_file_path}")
        
        try:
            # 读取txt文件
            with open(txt_file_path, 'r', encoding='utf-8') as f:
                content = f.read()
            
            # 创建新的Word文档
            doc = Document()
            
            # 分割内容为文本和HTML部分
            parts = re.split(r'(<[^>]+>.*?</[^>]+>|<[^>]+/>)', content, flags=re.DOTALL)
            
            for part in parts:
                if not part.strip():
                    continue
                
                # 检查是否为HTML内容
                if re.match(r'<[^>]+>', part.strip()):
                    # 这是HTML内容，需要转换为图片
                    print("发现HTML内容，正在转换为图片...")
                    
                    # 将HTML转换为图片
                    image_data = self.html_to_image(part.strip())
                    
                    if image_data:
                        # 添加图片到文档
                        self.add_image_to_docx(doc, image_data)
                        print("HTML图片已添加到文档")
                    else:
                        print("HTML转图片失败")
                else:
                    # 这是普通文本内容
                    clean_text = self.clean_text_content(part)
                    if clean_text:
                        # 添加文本到文档
                        doc.add_paragraph(clean_text)
            
            # 保存文档
            doc.save(output_docx_path)
            print(f"文档已保存: {output_docx_path}")
            return True
            
        except Exception as e:
            print(f"转换文件时出错: {e}")
            return False
    
    def batch_convert(self, input_dir, output_dir):
        """批量转换目录中的所有txt文件"""
        if not os.path.exists(output_dir):
            os.makedirs(output_dir)
        
        success_count = 0
        total_count = 0
        
        for filename in os.listdir(input_dir):
            if filename.endswith('.txt'):
                txt_path = os.path.join(input_dir, filename)
                docx_filename = filename.replace('.txt', '.docx')
                docx_path = os.path.join(output_dir, docx_filename)
                
                total_count += 1
                print(f"\n处理文件 {total_count}: {filename}")
                
                if self.convert_txt_to_docx(txt_path, docx_path):
                    success_count += 1
        
        print(f"\n转换完成！成功转换 {success_count}/{total_count} 个文件")
        return success_count, total_count

def main():
    """主函数"""
    converter = SimpleHTMLConverter()
    
    print("简单HTML转Docx转换器")
    print("=" * 50)
    
    # 检查输入目录
    input_dir = "generated_proposal"
    output_dir = "converted_docx"
    
    if not os.path.exists(input_dir):
        print(f"输入目录不存在: {input_dir}")
        return
    
    # 查找所有txt文件
    txt_files = [f for f in os.listdir(input_dir) if f.endswith('.txt')]
    
    if not txt_files:
        print("未找到txt文件")
        return
    
    print(f"找到 {len(txt_files)} 个txt文件")
    
    # 批量转换
    success, total = converter.batch_convert(input_dir, output_dir)
    
    if success > 0:
        print(f"\n转换成功！请查看输出目录: {output_dir}")
    else:
        print("\n转换失败")

if __name__ == "__main__":
    main() 