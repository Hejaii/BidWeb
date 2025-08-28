import os
import re
import requests
from docx import Document
from docx.shared import Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
import tempfile
from PIL import Image, ImageDraw, ImageFont
import io

class SimpleHTMLToDocxConverter:
    def __init__(self):
        # 通义千问 API Key
        self.api_key = "sk-fe0485c281964259b404907d483d3777"
        self.api_url = "https://dashscope.aliyuncs.com/api/v1/services/aigc/text-generation/generation"
        
    def call_qwen_api(self, prompt):
        """调用通义千问API"""
        headers = {
            "Authorization": f"Bearer {self.api_key}",
            "Content-Type": "application/json"
        }
        payload = {
            "model": "qwen-turbo",
            "input": {
                "messages": [
                    {
                        "role": "user",
                        "content": prompt
                    }
                ]
            }
        }
        response = requests.post(self.api_url, headers=headers, json=payload)
        return response.json()
    
    def create_technical_image(self, title, description):
        """创建技术架构图片"""
        try:
            # 创建图片
            width, height = 600, 400
            img = Image.new('RGB', (width, height), color='#f8f9fa')
            draw = ImageDraw.Draw(img)
            
            # 尝试加载中文字体
            try:
                font_large = ImageFont.truetype("arial.ttf", 24)
                font_medium = ImageFont.truetype("arial.ttf", 16)
                font_small = ImageFont.truetype("arial.ttf", 12)
            except:
                font_large = ImageFont.load_default()
                font_medium = ImageFont.load_default()
                font_small = ImageFont.load_default()
            
            # 绘制标题
            title_bbox = draw.textbbox((0, 0), title, font=font_large)
            title_width = title_bbox[2] - title_bbox[0]
            title_x = (width - title_width) // 2
            draw.text((title_x, 30), title, fill='#2c3e50', font=font_large)
            
            # 绘制描述
            lines = self.wrap_text(description, font_medium, width - 40)
            y_offset = 80
            for line in lines:
                line_bbox = draw.textbbox((0, 0), line, font=font_medium)
                line_width = line_bbox[2] - line_bbox[0]
                line_x = (width - line_width) // 2
                draw.text((line_x, y_offset), line, fill='#34495e', font=font_medium)
                y_offset += 25
            
            # 绘制技术架构框
            self.draw_architecture_boxes(draw, width, height, font_small)
            
            # 保存到内存
            img_buffer = io.BytesIO()
            img.save(img_buffer, format='PNG')
            img_buffer.seek(0)
            
            return img_buffer.getvalue()
            
        except Exception as e:
            print(f"创建技术图片时出错: {e}")
            return self.create_default_image()
    
    def wrap_text(self, text, font, max_width):
        """文本换行"""
        words = text.split()
        lines = []
        current_line = ""
        
        for word in words:
            test_line = current_line + " " + word if current_line else word
            bbox = ImageDraw.Draw(Image.new('RGB', (1, 1))).textbbox((0, 0), test_line, font=font)
            text_width = bbox[2] - bbox[0]
            
            if text_width <= max_width:
                current_line = test_line
            else:
                if current_line:
                    lines.append(current_line)
                current_line = word
        
        if current_line:
            lines.append(current_line)
        
        return lines
    
    def draw_architecture_boxes(self, draw, width, height, font):
        """绘制技术架构框"""
        # 绘制三个主要模块
        boxes = [
            {"x": 50, "y": 200, "w": 150, "h": 80, "text": "数据层", "color": "#3498db"},
            {"x": 225, "y": 200, "w": 150, "h": 80, "text": "服务层", "color": "#e74c3c"},
            {"x": 400, "y": 200, "w": 150, "h": 80, "text": "应用层", "color": "#2ecc71"}
        ]
        
        for box in boxes:
            # 绘制框
            draw.rectangle([box["x"], box["y"], box["x"] + box["w"], box["y"] + box["h"]], 
                         outline=box["color"], width=2, fill=box["color"] + "20")
            
            # 绘制文字
            text_bbox = draw.textbbox((0, 0), box["text"], font=font)
            text_width = text_bbox[2] - text_bbox[0]
            text_x = box["x"] + (box["w"] - text_width) // 2
            text_y = box["y"] + (box["h"] - 20) // 2
            draw.text((text_x, text_y), box["text"], fill='#2c3e50', font=font)
    
    def create_default_image(self):
        """创建默认图片"""
        try:
            img = Image.new('RGB', (600, 400), color='#ecf0f1')
            draw = ImageDraw.Draw(img)
            
            try:
                font = ImageFont.truetype("arial.ttf", 20)
            except:
                font = ImageFont.load_default()
            
            draw.text((200, 190), "技术架构图", fill='#7f8c8d', font=font)
            
            img_buffer = io.BytesIO()
            img.save(img_buffer, format='PNG')
            img_buffer.seek(0)
            
            return img_buffer.getvalue()
        except Exception as e:
            print(f"创建默认图片时出错: {e}")
            return b""
    
    def extract_image_descriptions(self, content):
        """提取图片描述"""
        # 匹配图片描述的正则表达式
        pattern = r'（此处需\[(.*?)\]的图片）'
        matches = re.findall(pattern, content)
        return matches
    
    def generate_image_from_description(self, description, context=""):
        """根据描述生成图片"""
        prompt = f"""请根据以下描述生成一个技术架构图的标题和简要说明：

描述：{description}
上下文：{context}

要求：
1. 生成一个简洁的技术架构图标题（10字以内）
2. 生成一个简要的技术说明（20字以内）
3. 返回格式：标题|说明

请直接返回结果，不要解释："""
        
        try:
            response = self.call_qwen_api(prompt)
            result = response['output']['text'].strip()
            
            if '|' in result:
                title, desc = result.split('|', 1)
                return self.create_technical_image(title.strip(), desc.strip())
            else:
                return self.create_technical_image("技术架构图", description)
                
        except Exception as e:
            print(f"生成图片描述时出错: {e}")
            return self.create_technical_image("技术架构图", description)
    
    def add_image_to_docx(self, doc, image_data, width_cm=14.0):
        """将图片添加到docx文档"""
        try:
            # 将图片数据写入临时文件
            with tempfile.NamedTemporaryFile(delete=False, suffix='.png') as temp_file:
                temp_file.write(image_data)
                temp_file_path = temp_file.name
            
            # 添加图片到文档
            paragraph = doc.add_paragraph()
            run = paragraph.add_run()
            run.add_picture(temp_file_path, width=Cm(width_cm))
            
            # 设置图片居中
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # 清理临时文件
            os.unlink(temp_file_path)
            
        except Exception as e:
            print(f"添加图片到文档时出错: {e}")
    
    def clean_text_content(self, text):
        """清理文本内容"""
        # 移除图片描述标记
        clean_text = re.sub(r'（此处需\[.*?\]的图片）', '', text)
        # 清理多余的空行
        clean_text = re.sub(r'\n\s*\n', '\n\n', clean_text)
        return clean_text.strip()
    
    def convert_txt_to_docx(self, txt_file_path, output_docx_path):
        """将带图片描述的txt文件转换为docx文件"""
        print(f"开始转换文件: {txt_file_path}")
        
        try:
            # 读取txt文件
            with open(txt_file_path, 'r', encoding='utf-8') as f:
                content = f.read()
            
            # 创建新的Word文档
            doc = Document()
            
            # 提取图片描述
            image_descriptions = self.extract_image_descriptions(content)
            
            # 分割内容
            parts = re.split(r'（此处需\[.*?\]的图片）', content)
            
            for i, part in enumerate(parts):
                if not part.strip():
                    continue
                
                # 添加文本内容
                clean_text = self.clean_text_content(part)
                if clean_text:
                    doc.add_paragraph(clean_text)
                
                # 如果有对应的图片描述，添加图片
                if i < len(image_descriptions):
                    print(f"正在生成图片: {image_descriptions[i]}")
                    image_data = self.generate_image_from_description(image_descriptions[i], part)
                    
                    if image_data:
                        self.add_image_to_docx(doc, image_data)
                        print("图片已添加到文档")
            
            # 保存文档
            doc.save(output_docx_path)
            print(f"文档已保存: {output_docx_path}")
            return True
            
        except Exception as e:
            print(f"转换文件时出错: {e}")
            return False
    
    def batch_convert(self, input_dir, output_dir):
        """批量转换目录中的所有txt文件"""
        if not os.path.exists(output_dir):
            os.makedirs(output_dir)
        
        success_count = 0
        total_count = 0
        
        for filename in os.listdir(input_dir):
            if filename.endswith('.txt'):
                txt_path = os.path.join(input_dir, filename)
                docx_filename = filename.replace('.txt', '.docx')
                docx_path = os.path.join(output_dir, docx_filename)
                
                total_count += 1
                print(f"\n处理文件 {total_count}: {filename}")
                
                if self.convert_txt_to_docx(txt_path, docx_path):
                    success_count += 1
        
        print(f"\n转换完成！成功转换 {success_count}/{total_count} 个文件")
        return success_count, total_count

def main():
    """主函数"""
    converter = SimpleHTMLToDocxConverter()
    
    print("TXT转Docx转换器（带图片生成）")
    print("=" * 50)
    
    # 检查输入目录
    input_dir = "generated_proposal"
    output_dir = "converted_docx"
    
    if not os.path.exists(input_dir):
        print(f"输入目录不存在: {input_dir}")
        return
    
    # 查找所有txt文件
    txt_files = [f for f in os.listdir(input_dir) if f.endswith('.txt')]
    
    if not txt_files:
        print("未找到txt文件")
        return
    
    print(f"找到 {len(txt_files)} 个txt文件")
    
    # 批量转换
    success, total = converter.batch_convert(input_dir, output_dir)
    
    if success > 0:
        print(f"\n转换成功！请查看输出目录: {output_dir}")
    else:
        print("\n转换失败")

if __name__ == "__main__":
    main() 