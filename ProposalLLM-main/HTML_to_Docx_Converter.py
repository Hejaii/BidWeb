import os
import re
import json
import requests
from docx import Document
from docx.shared import Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.shared import OxmlElement, qn
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
import base64
from PIL import Image
import io
from bs4 import BeautifulSoup
import html2image
import tempfile

class HTMLToDocxConverter:
    def __init__(self):
        # 通义千问 API Key
        self.api_key = "sk-fe0485c281964259b404907d483d3777"
        self.api_url = "https://dashscope.aliyuncs.com/api/v1/services/aigc/text-generation/generation"
        
        # 图片缓存
        self.image_cache = {}
        
    def call_qwen_api(self, prompt):
        """调用通义千问API"""
        headers = {
            "Authorization": f"Bearer {self.api_key}",
            "Content-Type": "application/json"
        }
        payload = {
            "model": "qwen-turbo",
            "input": {
                "messages": [
                    {
                        "role": "user",
                        "content": prompt
                    }
                ]
            }
        }
        response = requests.post(self.api_url, headers=headers, json=payload)
        return response.json()
    
    def extract_html_content(self, txt_file_path):
        """从txt文件中提取HTML内容"""
        try:
            with open(txt_file_path, 'r', encoding='utf-8') as f:
                content = f.read()
            
            # 查找HTML标签
            html_pattern = r'<[^>]+>.*?</[^>]+>|<[^>]+/>'
            html_matches = re.findall(html_pattern, content, re.DOTALL)
            
            return content, html_matches
        except Exception as e:
            print(f"读取文件时出错: {e}")
            return None, []
    
    def generate_image_from_html(self, html_content, context=""):
        """从HTML内容生成图片"""
        try:
            # 使用html2image将HTML转换为图片
            hti = html2image.Html2Image()
            
            # 创建完整的HTML文档
            full_html = f"""
            <!DOCTYPE html>
            <html>
            <head>
                <meta charset="UTF-8">
                <style>
                    body {{ 
                        font-family: 'Microsoft YaHei', Arial, sans-serif; 
                        margin: 0; 
                        padding: 20px; 
                        background: white; 
                    }}
                    .container {{ 
                        max-width: 600px; 
                        margin: 0 auto; 
                    }}
                </style>
            </head>
            <body>
                <div class="container">
                    {html_content}
                </div>
            </body>
            </html>
            """
            
            # 生成图片
            output_path = hti.html2image(html_str=full_html, save_as='temp_image.png')[0]
            
            # 读取生成的图片
            with open(output_path, 'rb') as f:
                image_data = f.read()
            
            # 清理临时文件
            os.remove(output_path)
            
            return image_data
            
        except Exception as e:
            print(f"生成图片时出错: {e}")
            # 返回默认图片
            return self.create_default_image()
    
    def create_default_image(self):
        """创建默认图片"""
        try:
            # 创建一个简单的默认图片
            img = Image.new('RGB', (600, 400), color='#f0f0f0')
            
            # 保存到内存
            img_buffer = io.BytesIO()
            img.save(img_buffer, format='PNG')
            img_buffer.seek(0)
            
            return img_buffer.getvalue()
        except Exception as e:
            print(f"创建默认图片时出错: {e}")
            return b""
    
    def add_image_to_docx(self, doc, image_data, width_cm=14.0):
        """将图片添加到docx文档"""
        try:
            # 将图片数据写入临时文件
            with tempfile.NamedTemporaryFile(delete=False, suffix='.png') as temp_file:
                temp_file.write(image_data)
                temp_file_path = temp_file.name
            
            # 添加图片到文档
            paragraph = doc.add_paragraph()
            run = paragraph.add_run()
            run.add_picture(temp_file_path, width=Cm(width_cm))
            
            # 设置图片居中
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # 清理临时文件
            os.unlink(temp_file_path)
            
        except Exception as e:
            print(f"添加图片到文档时出错: {e}")
    
    def clean_text_content(self, text):
        """清理文本内容，移除HTML标签"""
        # 移除HTML标签
        clean_text = re.sub(r'<[^>]+>', '', text)
        # 清理多余的空行
        clean_text = re.sub(r'\n\s*\n', '\n\n', clean_text)
        return clean_text.strip()
    
    def convert_html_to_docx(self, txt_file_path, output_docx_path):
        """将带HTML的txt文件转换为docx文件"""
        print(f"开始转换文件: {txt_file_path}")
        
        # 创建新的Word文档
        doc = Document()
        
        # 提取HTML内容
        content, html_matches = self.extract_html_content(txt_file_path)
        
        if not content:
            print("无法读取文件内容")
            return False
        
        # 分割内容为文本和HTML部分
        parts = re.split(r'(<[^>]+>.*?</[^>]+>|<[^>]+/>)', content, flags=re.DOTALL)
        
        for part in parts:
            if not part.strip():
                continue
            
            # 检查是否为HTML内容
            if re.match(r'<[^>]+>', part.strip()):
                # 这是HTML内容，需要转换为图片
                print("发现HTML内容，正在生成图片...")
                
                # 生成图片
                image_data = self.generate_image_from_html(part.strip())
                
                if image_data:
                    # 添加图片到文档
                    self.add_image_to_docx(doc, image_data)
                    print("图片已添加到文档")
                else:
                    print("图片生成失败")
            else:
                # 这是普通文本内容
                clean_text = self.clean_text_content(part)
                if clean_text:
                    # 添加文本到文档
                    doc.add_paragraph(clean_text)
        
        # 保存文档
        try:
            doc.save(output_docx_path)
            print(f"文档已保存: {output_docx_path}")
            return True
        except Exception as e:
            print(f"保存文档时出错: {e}")
            return False
    
    def batch_convert(self, input_dir, output_dir):
        """批量转换目录中的所有txt文件"""
        if not os.path.exists(output_dir):
            os.makedirs(output_dir)
        
        success_count = 0
        total_count = 0
        
        for filename in os.listdir(input_dir):
            if filename.endswith('.txt'):
                txt_path = os.path.join(input_dir, filename)
                docx_filename = filename.replace('.txt', '.docx')
                docx_path = os.path.join(output_dir, docx_filename)
                
                total_count += 1
                print(f"\n处理文件 {total_count}: {filename}")
                
                if self.convert_html_to_docx(txt_path, docx_path):
                    success_count += 1
        
        print(f"\n转换完成！成功转换 {success_count}/{total_count} 个文件")
        return success_count, total_count

def main():
    """主函数"""
    converter = HTMLToDocxConverter()
    
    print("HTML转Docx转换器")
    print("=" * 50)
    
    # 检查输入目录
    input_dir = "generated_proposal"
    output_dir = "converted_docx"
    
    if not os.path.exists(input_dir):
        print(f"输入目录不存在: {input_dir}")
        return
    
    # 查找所有txt文件
    txt_files = [f for f in os.listdir(input_dir) if f.endswith('.txt')]
    
    if not txt_files:
        print("未找到txt文件")
        return
    
    print(f"找到 {len(txt_files)} 个txt文件")
    
    # 批量转换
    success, total = converter.batch_convert(input_dir, output_dir)
    
    if success > 0:
        print(f"\n转换成功！请查看输出目录: {output_dir}")
    else:
        print("\n转换失败")

if __name__ == "__main__":
    main() 